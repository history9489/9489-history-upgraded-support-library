#************* Develop on the 14th Aug 2026 *****************
# ********** 9696 Geography PYP Portal (Custom Color Theme) 9th Aug 2026 ***********
import datetime
import io
import os
import fitz  # PyMuPDF
import streamlit as st

# Word Document Libraries
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Google API Libraries
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload

# ==========================================
# 0. STREAMLIT PAGE CONFIG & CUSTOM STYLING
# ==========================================
st.set_page_config(
    page_title="9696 Geography PYP Portal", 
    page_icon="🌍",
    layout="wide"
)

# Custom CSS Theme Injecting User Specified Color Palette
st.markdown("""
    <style>
    /* 1. Main Page Background */
    .stApp, [data-testid="stAppViewContainer"] {
        background-color: #63D0F8 !important;
    }
    
    /* 2. Sidebar Background */
    [data-testid="stSidebar"], [data-testid="stSidebar"] > div:first-child {
        background-color: #F0FCBB !important;
    }

    /* 3. Global Text Color (#384403) */
    html, body, [class*="css"], h1, h2, h3, h4, h5, h6, p, span, label, div, .stMarkdown {
        color: #384403 !important;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }

    /* 4. Input Bars, Selectboxes, Text Areas & Borders (#A5C809) */
    div[data-baseweb="input"], 
    div[data-baseweb="select"] > div, 
    .stTextInput input, 
    .stSelectbox select,
    textarea {
        background-color: #CEC2F5 !important;
        color: #141801 !important;
        border-radius: 10px !important;
        border: 5px solid #A5C809 !important;
    }

    /* 5. Buttons Styling (#C9F40B with #A5C809 Border) */
    .stButton button, 
    .stDownloadButton button, 
    [data-testid="baseButton-secondary"], 
    [data-testid="baseButton-primary"] {
        background-color: #C9F40B !important;
        color: #384403 !important;
        border: 2.3px solid #A5C809 !important;
        border-radius: 8px !important;
        font-weight: bold !important;
        transition: all 0.2s ease-in-out;
    }
    
    /* Hover state for buttons */
    .stButton button:hover, .stDownloadButton button:hover {
        background-color: #A5C809 !important;
        color: #384403 !important;
        border: 4px solid #384403 !important;
    }

    /* 6. Navigation Tab Labels */
    button[data-baseweb="tab"] p {
        font-weight: bold !important;
        font-size: 2.05rem !important;
        color: #384403 !important;
    }
    
    /* Active Tab Highlight Indicator */
    div[data-baseweb="tab-highlight"] {
        background-color: #F863E1 !important;
    }

    /* 7. Expanders & Containers Border Color */
    [data-testid="stExpander"] {
        border: 1.5px solid #A5C809 !important;
        border-radius: 8px !important;
        background-color: #F0FCBB !important;
    }
    </style>
""", unsafe_allow_html=True)


# ==========================================
# 1. DIRECTORY MAPPING & CONFIGURATION
# ==========================================
SYLLABUS_CODE = "9696"
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']

LOCAL_FOLDERS = {
    "p1_physical": "core_physicalP1",
    "p2_human": "core_humanP2",
    "p3_adv_physical": "Adv_physicalP3",
    "p4_adv_human": "Adv_humanP4",
    "insert_p1_p3": "Insertpyp_P1P3",
    "insert_p2_p4": "Insertpyp_P2P4",
    "ms_p1_p2": "marksch_P1P2",
    "ms_p3_p4": "marksch_P3P4"
}

for folder_path in LOCAL_FOLDERS.values():
    os.makedirs(folder_path, exist_ok=True)

# ==========================================
# 2. SERVICE ACCOUNT AUTHENTICATION & SYNC
# ==========================================
def build_drive_service(write_access=False):
    """Authenticates using Google Service Account credentials from Streamlit Secrets."""
    try:
        if "gcp_service_account" in st.secrets:
            service_account_info = dict(st.secrets["gcp_service_account"])
            scopes = ['https://www.googleapis.com/auth/drive.file'] if write_access else SCOPES
            creds = service_account.Credentials.from_service_account_info(
                service_account_info, 
                scopes=scopes
            )
            return build('drive', 'v3', credentials=creds)
        else:
            st.error("Missing [gcp_service_account] configuration in secrets.")
            return None
    except Exception as e:
        st.error(f"Authentication Error: {e}")
        return None

def sync_drive_folder_to_local(folder_key: str) -> tuple[int, str]:
    """Downloads missing files from Google Drive folder into local directory."""
    service = build_drive_service(write_access=False)
    if not service:
        return 0, "Failed to authenticate Service Account."
    
    folder_ids = st.secrets.get("drive_folders", {})
    drive_folder_id = folder_ids.get(folder_key)
    
    if not drive_folder_id:
        return 0, f"Missing drive_folder_id for `{folder_key}` in secrets."

    local_path = LOCAL_FOLDERS[folder_key]
    
    try:
        query = f"'{drive_folder_id}' in parents and trashed = false"
        drive_files = []
        page_token = None

        while True:
            response = service.files().list(
                q=query,
                fields="nextPageToken, files(id, name, mimeType)",
                pageToken=page_token,
                pageSize=100
            ).execute()
            
            drive_files.extend(response.get('files', []))
            page_token = response.get('nextPageToken', None)
            
            if not page_token:
                break

        downloaded_count = 0

        for file_info in drive_files:
            file_name = file_info['name']
            file_id = file_info['id']
            local_file_path = os.path.join(local_path, file_name)

            if not os.path.exists(local_file_path):
                request = service.files().get_media(fileId=file_id)
                with open(local_file_path, "wb") as f:
                    downloader = MediaIoBaseDownload(f, request)
                    done = False
                    while not done:
                        _, done = downloader.next_chunk()
                downloaded_count += 1

        total_local_files = len([f for f in os.listdir(local_path) if os.path.isfile(os.path.join(local_path, f))])
        return downloaded_count, f"Synced {downloaded_count} new file(s) for `{folder_key}` (Total: {total_local_files})."
        
    except Exception as e:
        return 0, f"Sync error for `{folder_key}`: {e}"

def perform_bulk_sync():
    """Syncs all 8 configured Google Drive folders."""
    total_synced = 0
    messages = []
    for f_key in LOCAL_FOLDERS.keys():
        count, msg = sync_drive_folder_to_local(f_key)
        total_synced += count
        messages.append(msg)
    return total_synced, messages

# ==========================================
# 3. HELPER FUNCTIONS
# ==========================================
def create_worksheet_docx(basket_items: list) -> io.BytesIO:
    """Generates a Word document containing selected PDF pages."""
    doc = Document()
    section = doc.sections[0]

    section.page_width = Inches(8.5)
    section.page_height = Inches(11.5)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run("Page ")
    add_page_number_to_run(header_run)

    doc.add_heading(f'PTES {SYLLABUS_CODE} Geography Worksheet', level=1)

    for idx, item in enumerate(basket_items):
        doc.add_heading(f"Source: {item['file']} (Page {item['page'] + 1})", level=2)
        pdf_doc = fitz.open(item['path'])
        page = pdf_doc.load_page(item['page'])
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_data = io.BytesIO(pix.tobytes("png"))

        doc.add_picture(img_data, width=Inches(7.8), height=Inches(8.8))

        if idx < len(basket_items) - 1:
            doc.add_page_break()
        pdf_doc.close()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def add_page_number_to_run(run):
    """Adds a dynamic Word page number field."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def render_pdf_page_preview(filepath: str, page_num: int = 0):
    """Renders a PDF page to PNG image bytes for preview."""
    try:
        doc = fitz.open(filepath)
        page = doc.load_page(page_num)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
        img_bytes = pix.tobytes("png")
        doc.close()
        return img_bytes
    except Exception as e:
        st.error(f"Unable to render page preview: {e}")
        return None

def execute_pdf_search(folder_key: str, keyword_string: str) -> list[dict]:
    """Searches PDF files in a specific folder for matching keywords."""
    results = []
    keywords = [k.strip().lower() for k in keyword_string.split(",") if k.strip()]
    folder_path = LOCAL_FOLDERS[folder_key]
    
    if os.path.exists(folder_path):
        for file in os.listdir(folder_path):
            if file.endswith(".pdf"):
                filepath = os.path.join(folder_path, file)
                try:
                    doc = fitz.open(filepath)
                    for page_num in range(len(doc)):
                        text = doc[page_num].get_text().lower()
                        if all(kw in text for kw in keywords):
                            results.append({
                                "file": file, 
                                "page": page_num, 
                                "path": filepath
                            })
                    doc.close()
                except Exception:
                    continue
    return results

# ==========================================
# 4. SESSION STATE INITIALIZATION
# ==========================================
if 'handout_basket' not in st.session_state:
    st.session_state.handout_basket = []

if 'p1_p3_results' not in st.session_state:
    st.session_state.p1_p3_results = []
if 'p2_p4_results' not in st.session_state:
    st.session_state.p2_p4_results = []
if 'insert_results' not in st.session_state:
    st.session_state.insert_results = []

if 'has_auto_synced' not in st.session_state:
    st.session_state.has_auto_synced = True
    with st.spinner("🚀 Waking up portal & auto-syncing Geography files via Service Account..."):
        perform_bulk_sync()

# ==========================================
# 5. STREAMLIT UI LAYOUT
# ==========================================
st.title("PUSAT TINGKATAN ENAM SENGKURONG")
st.subheader(f"🌍 {SYLLABUS_CODE} Geography PYP Resource Library")

# --- SIDEBAR CONTROLS ---
with st.sidebar:
    st.header("🔄 Google Drive Sync")
    if st.button("🔄 Sync Google Drive", use_container_width=True):
        with st.spinner("Syncing Google Drive folders..."):
            synced_count, sync_msgs = perform_bulk_sync()
            st.success(f"Sync Complete! {synced_count} new file(s) downloaded.")
            for m in sync_msgs:
                st.caption(m)

    st.markdown("---")
    st.metric(label="Saved Pages in Basket", value=len(st.session_state.handout_basket))

    if st.button("🗑️ Clear Basket", use_container_width=True):
        st.session_state.handout_basket = []
        st.rerun()

# --- NAVIGATION TABS ---
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "🏔️ Physical Geog", 
    "🏙️ Human Geog", 
    "🗺️ PYP Inserts ",
    "🔑 PYP Answer", 
    "🛒 PYP Cart", 
    "⚙️ Upload PYP/Admin"
])

# --- TAB 1: PHYSICAL GEOGRAPHY (PAPER 1 or PAPER 3) ---
with tab1:
    st.subheader("🏔️ Physical Geography Search (Paper 1 / Paper 3)")
    
    selected_phys_paper = st.selectbox(
        "Select Component Paper:", 
        options=["p1_physical", "p3_adv_physical"],
        format_func=lambda x: "Paper 1 (12): core_physicalP1" if x == "p1_physical" else "Paper 3 (32): Adv_physicalP3",
        key="select_phys_paper"
    )

    t_kw = st.text_input(
        "Enter Keywords", 
        placeholder="e.g., hydrology, fluvial, atmosphere, weathering, coastal, hazardous", 
        key="phys_kw"
    )

    if st.button("Search Keyword", key="btn_search_phys"):
        if t_kw.strip():
            with st.spinner("Scanning PDFs..."):
                st.session_state.p1_p3_results = execute_pdf_search(selected_phys_paper, t_kw)
        else:
            st.warning("Please enter a keyword.")

    if st.session_state.p1_p3_results:
        st.write(f"Found **{len(st.session_state.p1_p3_results)}** matching page(s):")
        for idx, item in enumerate(st.session_state.p1_p3_results):
            with st.expander(f"📄 {item['file']} | Page {item['page'] + 1}"):
                c1, c2 = st.columns([3, 1])
                with c1:
                    preview_img = render_pdf_page_preview(item["path"], item["page"])
                    if preview_img:
                        st.image(preview_img, use_container_width=True)
                with c2:
                    if st.button("➕ Add to Cart", key=f"add_phys_{idx}"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added Page {item['page'] + 1} to basket!")
                        st.rerun()
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    with open(item["path"], "rb") as pdf_f:
                        st.download_button(
                            label="📥 Download Full PDF",
                            data=pdf_f,
                            file_name=item["file"],
                            mime="application/pdf",
                            key=f"dl_phys_{idx}"
                        )

# --- TAB 2: HUMAN GEOGRAPHY (PAPER 2 or PAPER 4) ---
with tab2:
    st.subheader("🏙️ Human Geography Search (Paper 2 / Paper 4)")
    
    selected_human_paper = st.selectbox(
        "Select Component Paper:", 
        options=["p2_human", "p4_adv_human"],
        format_func=lambda x: "Paper 2 (22): core_humanP2" if x == "p2_human" else "Paper 4 (42): Adv_humanP4",
        key="select_human_paper"
    )

    h_kw = st.text_input(
        "Enter Keywords", 
        placeholder="e.g., population, migration, urban, settlement, energy, governance", 
        key="human_kw"
    )

    if st.button("Search Keyword", key="btn_search_human"):
        if h_kw.strip():
            with st.spinner("Scanning PDFs..."):
                st.session_state.p2_p4_results = execute_pdf_search(selected_human_paper, h_kw)
        else:
            st.warning("Please enter a keyword.")

    if st.session_state.p2_p4_results:
        st.write(f"Found **{len(st.session_state.p2_p4_results)}** matching page(s):")
        for idx, item in enumerate(st.session_state.p2_p4_results):
            with st.expander(f"📄 {item['file']} | Page {item['page'] + 1}"):
                c1, c2 = st.columns([3, 1])
                with c1:
                    preview_img = render_pdf_page_preview(item["path"], item["page"])
                    if preview_img:
                        st.image(preview_img, use_container_width=True)
                with c2:
                    if st.button("➕ Add to Cart", key=f"add_human_{idx}"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added Page {item['page'] + 1} to basket!")
                        st.rerun()
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    with open(item["path"], "rb") as pdf_f:
                        st.download_button(
                            label="📥 Download Full PDF",
                            data=pdf_f,
                            file_name=item["file"],
                            mime="application/pdf",
                            key=f"dl_human_{idx}"
                        )

# --- TAB 3: INSERTS SEARCH (P1, P2, P3, P4) ---
with tab3:
    st.subheader("🗺️ Geography Inserts Search (Figures, Maps & Tables)")
    
    selected_insert = st.selectbox(
        "Select Insert Paper Target:", 
        options=["insert_p1", "insert_p2", "insert_p3", "insert_p4"],
        format_func=lambda x: {
            "insert_p1": "Insert Materials from Paper 1",
            "insert_p2": "Insert Materials from Paper 2",
            "insert_p3": "Insert Materials from Paper 3",
            "insert_p4": "Insert Materials from Paper 4"
        }[x],
        key="select_insert_target"
    )

    ins_kw = st.text_input(
        "Enter Keywords or Fig/Table number", 
        placeholder="e.g., Fig 1.1, Table 2, hydrograph, satellite image", 
        key="insert_kw"
    )

    folder_mapping = {
        "insert_p1": "insert_p1_p3",
        "insert_p3": "insert_p1_p3",
        "insert_p2": "insert_p2_p4",
        "insert_p4": "insert_p2_p4"
    }
    target_folder_key = folder_mapping[selected_insert]

    if st.button("Search Insert", key="btn_search_insert"):
        if ins_kw.strip():
            with st.spinner("Scanning Inserts..."):
                st.session_state.insert_results = execute_pdf_search(target_folder_key, ins_kw)
        else:
            st.warning("Please enter a search keyword.")

    if st.session_state.insert_results:
        st.write(f"Found **{len(st.session_state.insert_results)}** matching insert page(s):")
        for idx, item in enumerate(st.session_state.insert_results):
            with st.expander(f"📄 {item['file']} | Page {item['page'] + 1}"):
                c1, c2 = st.columns([3, 1])
                with c1:
                    preview_img = render_pdf_page_preview(item["path"], item["page"])
                    if preview_img:
                        st.image(preview_img, use_container_width=True)
                with c2:
                    if st.button("➕ Add to Cart", key=f"add_ins_{idx}"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added Page {item['page'] + 1} to basket!")
                        st.rerun()
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    with open(item["path"], "rb") as pdf_f:
                        st.download_button(
                            label="📥 Download Full Insert PDF",
                            data=pdf_f,
                            file_name=item["file"],
                            mime="application/pdf",
                            key=f"dl_ins_{idx}"
                        )

# --- TAB 4: DOWNLOAD MARK SCHEME ---
with tab4:
    st.subheader("🔑 Download Answer / Mark Schemes")
    st.caption("Select the exam session parameters or search your local mark scheme repositories.")
    
    col_y, col_m, col_v = st.columns([1, 2, 2])
    with col_y:
        as_year = st.text_input(
            "Year", 
            value="2023", 
            placeholder="e.g. 2023 or 23", 
            key="as_pyp_year"
        )
    with col_m:
        as_month = st.selectbox(
            "Select Session", 
            ["November / Autumn (w)", "June / Summer (s)", "March (m)"], 
            key="as_mth"
        )
        if "June" in as_month:
            month_code = "s"
        elif "March" in as_month:
            month_code = "m"
        else:
            month_code = "w"
            
    with col_v:
        as_variant = st.selectbox(
            "Select Component Variant", 
            ["11", "12", "13", "21", "22", "23", "31", "32", "33", "41", "42", "43"], 
            index=2, # Default to 13
            key="as_var"
        )

    cleaned_year = as_year.strip()
    short_year = cleaned_year[-2:] if len(cleaned_year) >= 2 else cleaned_year

    search_session_tag = f"{month_code}{short_year}"
    expected_ms_filename = f"{SYLLABUS_CODE}_{search_session_tag}_ms_{as_variant}.pdf"

    st.markdown("---")
    
    found_ms_files = []
    target_ms_folders = ["ms_p1_p2", "ms_p3_p4"]

    for folder_key in target_ms_folders:
        folder_path = LOCAL_FOLDERS[folder_key]
        if os.path.exists(folder_path):
            for file in os.listdir(folder_path):
                if file.endswith(".pdf"):
                    file_lower = file.lower()
                    if search_session_tag in file_lower and f"ms_{as_variant}" in file_lower:
                        found_ms_files.append(os.path.join(folder_path, file))
                    elif search_session_tag in file_lower and "ms" in file_lower and as_variant in file_lower:
                        if os.path.join(folder_path, file) not in found_ms_files:
                            found_ms_files.append(os.path.join(folder_path, file))

    if found_ms_files:
        st.success(f"Found {len(found_ms_files)} matching Answer Scheme file(s):")
        
        for ms_path in found_ms_files:
            ms_filename = os.path.basename(ms_path)
            
            with st.expander(f"🔑 Mark Scheme: {ms_filename}", expanded=True):
                col_dl, col_blank = st.columns([1, 2])
                with col_dl:
                    with open(ms_path, "rb") as f:
                        st.download_button(
                            label=f"📥 Download {ms_filename}",
                            data=f,
                            file_name=ms_filename,
                            mime="application/pdf",
                            key=f"dl_ms_btn_{ms_filename}"
                        )
                
                doc = fitz.open(ms_path)
                st.caption(f"Document contains {len(doc)} page(s). Previewing Page 1:")
                img_data = render_pdf_page_preview(ms_path, 0)
                if img_data:
                    st.image(img_data, caption=f"Preview: {ms_filename} (Page 1)", use_container_width=True)
                doc.close()
    else:
        st.warning(f"No Mark Scheme found matching session `{search_session_tag}` and variant `{as_variant}` (Expected pattern: `{expected_ms_filename}`).")
        st.info("💡 **Tip**: Ensure you have run **Sync Google Drive** from the sidebar so the latest mark scheme PDFs are downloaded from Google Drive into `marksch_P1P2` or `marksch_P3P4`.")

# --- TAB 5: DOWNLOAD HANDOUT MERGED (CART) ---
with tab5:
    st.subheader("🛒 PYP Cart & Worksheet Generator")
    
    if len(st.session_state.handout_basket) > 0:
        st.info(f"### 📋 Review Selected Pages ({len(st.session_state.handout_basket)} items)")
        st.warning("Review your selected pages below. You can preview them or remove specific items before generating your Word worksheet.")
        st.markdown("---")
        
        items_to_display = list(st.session_state.handout_basket)
        
        for idx, item in enumerate(items_to_display):
            filename = item.get("file", "Unknown File")
            page_num = item.get("page", 0) + 1
            file_path = item.get("path", "")
            
            with st.expander(f"📄 Item #{idx + 1}: {filename} (Page {page_num})", expanded=True):
                col_preview, col_action = st.columns([3, 1])
                
                with col_preview:
                    if os.path.exists(file_path):
                        img_bytes = render_pdf_page_preview(file_path, item.get("page", 0))
                        if img_bytes:
                            st.image(img_bytes, caption=f"Preview: Page {page_num}", use_container_width=True)
                    else:
                        st.caption(f"Source file path: `{file_path}`")
                
                with col_action:
                    st.markdown("#### Actions")
                    remove_key = f"remove_btn_cart_item_{idx}_{filename}_{page_num}"
                    if st.button("🗑️ Remove Item", key=remove_key, use_container_width=True):
                        st.session_state.handout_basket.pop(idx)
                        st.toast(f"Removed item #{idx + 1} from cart!")
                        st.rerun()

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("---")
        st.markdown("### 📝 Export Worksheet")
        
        with st.spinner("Generating Word Worksheet..."):
            doc_buffer = create_worksheet_docx(st.session_state.handout_basket)
            target_filename = f"{SYLLABUS_CODE}_Geography_Worksheet.docx"

        st.download_button(
            label="🪄 Download Merged Word Document Worksheet",
            data=doc_buffer,
            file_name=target_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.info("🛒 Your PYP Cart is currently empty! Navigate to Physical Geog, Human Geog, or PYP Inserts tabs and click '➕ Add to Cart' to add questions here.")

# --- TAB 6: UPLOAD PYP / ADMIN DASHBOARD ---
with tab6:
    st.subheader("⚙️ Upload PYP / Admin Dashboard")
    st.caption("Secure admin access to your 8 Google Drive repositories inside `PYPMaterials_Geo9696`.")

    admin_pwd = st.secrets.get("ADMIN_PASSWORD", "")
    pwd_input = st.text_input("Enter Admin Password", type="password", key="admin_pwd_input")

    if pwd_input and pwd_input == admin_pwd:
        st.success("Authenticated as Administrator")
        st.markdown("---")
        
        st.markdown("### 🌐 Google Drive Web Repositories")
        st.info("💡 **Instructions**: Click any button below to open that specific folder in Google Drive. You can drag and drop your new PDF files directly into the web browser!")
        
        drive_links = st.secrets.get("drive_web_links", {})

        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.link_button("📘 core_physicalP1", drive_links.get("p1_physical", "https://drive.google.com"), use_container_width=True)
            st.markdown("<br>", unsafe_allow_html=True)
            st.link_button("🗺️ Insertpyp_P1P3", drive_links.get("insert_p1_p3", "https://drive.google.com"), use_container_width=True)
            
        with c2:
            st.link_button("📗 core_humanP2", drive_links.get("p2_human", "https://drive.google.com"), use_container_width=True)
            st.markdown("<br>", unsafe_allow_html=True)
            st.link_button("🗺️ Insertpyp_P2P4", drive_links.get("insert_p2_p4", "https://drive.google.com"), use_container_width=True)
            
        with c3:
            st.link_button("📙 Adv_physicalP3", drive_links.get("p3_adv_physical", "https://drive.google.com"), use_container_width=True)
            st.markdown("<br>", unsafe_allow_html=True)
            st.link_button("🔑 marksch_P1P2", drive_links.get("ms_p1_p2", "https://drive.google.com"), use_container_width=True)
            
        with c4:
            st.link_button("📕 Adv_humanP4", drive_links.get("p4_adv_human", "https://drive.google.com"), use_container_width=True)
            st.markdown("<br>", unsafe_allow_html=True)
            st.link_button("🔑 marksch_P3P4", drive_links.get("ms_p3_p4", "https://drive.google.com"), use_container_width=True)

    elif pwd_input:
        st.error("Incorrect Admin Password.")

# ==========================================
# 6. PORTAL FOOTER
# ==========================================
st.markdown("---")
SCHOOL_NAME = "Pusat Tingkatan Enam Sengkurong (PTES)"
SCHOOL_VISION = "Nurturing Resilient Leaders & Future-Ready Citizens"
DEVELOPER_NAME = "Cikgu Haziqah / Computer Science Department"

footer_html = f"""
<div style="text-align: center; padding: 15px 0px; font-family: sans-serif;">
    <p style="margin: 0; font-size: 1.0em; font-weight: bold; color: #384403;">🏫 {SCHOOL_NAME}</p>
    <p style="margin: 5px 0; font-size: 0.9em; font-style: italic; color: #384403;">"{SCHOOL_VISION}"</p>
    <p style="margin: 5px 0 0 0; font-size: 0.85em; font-weight: 600; color: #384403;">💻 Developed by {DEVELOPER_NAME}</p>
</div>
"""
st.markdown(footer_html, unsafe_allow_html=True)
